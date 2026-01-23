"""
Script to convert .docx file to readable markdown format
"""
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("python-docx library is required. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document


def docx_to_markdown(docx_path, output_path=None):
    """
    Convert a .docx file to markdown format
    
    Args:
        docx_path: Path to the .docx file
        output_path: Path for output file (optional, defaults to same name with .md extension)
    """
    docx_path = Path(docx_path)
    
    if not docx_path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")
    
    if output_path is None:
        output_path = docx_path.with_suffix('.md')
    else:
        output_path = Path(output_path)
    
    # Open the document
    doc = Document(docx_path)
    
    # Extract text from all paragraphs
    markdown_lines = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            # Check if it's a heading (bold or larger font)
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.replace('Heading ', '')
                try:
                    level_num = int(level)
                    markdown_lines.append('#' * level_num + ' ' + text)
                except ValueError:
                    markdown_lines.append('## ' + text)
            else:
                # Check for bold text
                is_bold = False
                for run in paragraph.runs:
                    if run.bold:
                        is_bold = True
                        break
                
                if is_bold:
                    markdown_lines.append('**' + text + '**')
                else:
                    markdown_lines.append(text)
            
            markdown_lines.append('')  # Add blank line between paragraphs
    
    # Also extract text from tables
    for table in doc.tables:
        markdown_lines.append('')  # Blank line before table
        for i, row in enumerate(table.rows):
            row_text = '| ' + ' | '.join([cell.text.strip() for cell in row.cells]) + ' |'
            markdown_lines.append(row_text)
            if i == 0:  # Add header separator after first row
                markdown_lines.append('| ' + ' | '.join(['---'] * len(row.cells)) + ' |')
        markdown_lines.append('')  # Blank line after table
    
    # Write to output file
    output_path.write_text('\n'.join(markdown_lines), encoding='utf-8')
    print(f"Successfully converted {docx_path.name} to {output_path.name}")
    return output_path


if __name__ == "__main__":
    # Default input file
    input_file = Path(__file__).parent / "Moghimi_ABT_HYD_182_midterm_sample_questions_with_answers.docx"
    
    if len(sys.argv) > 1:
        input_file = Path(sys.argv[1])
    
    output_file = docx_to_markdown(input_file)
    print(f"\nOutput saved to: {output_file}")
    print(f"\nFirst 500 characters of converted content:")
    print("-" * 50)
    print(output_file.read_text(encoding='utf-8')[:500])
