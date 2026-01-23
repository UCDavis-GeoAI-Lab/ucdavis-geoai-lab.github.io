"""
Script to convert markdown files to .docx format with enhanced formatting
"""
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx library is required. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn


def add_border_to_paragraph(paragraph):
    """Add a border around a paragraph to create a box effect"""
    pPr = paragraph._element.get_or_add_pPr()
    
    # Create border elements
    borders = OxmlElement('w:pBdr')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    
    pPr.append(borders)
    
    # Add shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)


def add_formatted_run(paragraph, text):
    """Add text with markdown formatting (bold, code, checkmark) - fixed version"""
    if not text:
        return
    
    # Check if checkmark exists and remove it from text to avoid duplication
    has_checkmark = '✓' in text
    # Remove checkmark from anywhere in the text
    text_clean = re.sub(r'✓\s*', '', text).strip()
    text_clean = re.sub(r'\s*✓', '', text_clean).strip()
    
    # Split by markdown formatting patterns - handle nested cases
    # Pattern: **bold** (which may contain checkmark and code), `code`, or regular text
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`)'
    parts = re.split(pattern, text_clean)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            # Bold text - may contain inline code, need to process it
            bold_content = part[2:-2].replace('✓', '').strip()
            
            # Check if bold text contains inline code
            if '`' in bold_content:
                # Process bold text that contains code
                code_pattern = r'(`[^`]+`)'
                bold_parts = re.split(code_pattern, bold_content)
                
                for bold_part in bold_parts:
                    if bold_part.startswith('`') and bold_part.endswith('`'):
                        # Code within bold - extract code text (remove backticks)
                        code_text = bold_part[1:-1]
                        run = paragraph.add_run(code_text)
                        run.bold = True
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(196, 49, 16)  # Dark red for code
                    else:
                        # Regular bold text
                        if bold_part.strip():
                            run = paragraph.add_run(bold_part)
                            run.bold = True
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
            else:
                # Plain bold text
                if bold_content:
                    run = paragraph.add_run(bold_content)
                    run.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
        elif part.startswith('`') and part.endswith('`'):
            # Inline code (not in bold) - remove backticks
            code_text = part[1:-1]
            run = paragraph.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(196, 49, 16)  # Dark red for code
        else:
            # Regular text - check for inline code within
            if part.strip():
                code_parts = re.split(r'(`[^`]+`)', part)
                for code_part in code_parts:
                    if code_part.startswith('`') and code_part.endswith('`'):
                        # Remove backticks from code
                        code_text = code_part[1:-1]
                        run = paragraph.add_run(code_text)
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(196, 49, 16)
                    else:
                        # Remove any checkmarks from regular text
                        clean_part = code_part.replace('✓', '').strip()
                        if clean_part:
                            run = paragraph.add_run(clean_part)
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
    
    # Add checkmark only once at the end if it was in the original text
    if has_checkmark:
        run = paragraph.add_run(' ✓')
        run.bold = True
        run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
        run.font.name = 'Arial'
        run.font.size = Pt(11)


def convert_markdown_simple(md_path, docx_path=None):
    """
    Enhanced conversion with better formatting
    """
    md_path = Path(md_path)
    
    if not md_path.exists():
        raise FileNotFoundError(f"File not found: {md_path}")
    
    if docx_path is None:
        docx_path = md_path.with_suffix('.docx')
    else:
        docx_path = Path(docx_path)
    
    # Read markdown file
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font to Arial (sans-serif)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Set heading fonts to Arial as well
    for heading_level in range(1, 10):
        try:
            heading_style = doc.styles[f'Heading {heading_level}']
            heading_style.font.name = 'Arial'
        except:
            pass
    
    lines = content.split('\n')
    in_code_block = False
    code_lines = []
    
    for i, line in enumerate(lines):
        stripped = line.rstrip()
        
        # Code block detection
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                continue
            else:
                # End of code block - create boxed paragraph
                in_code_block = False
                if code_lines:
                    # Join code lines
                    code_text = '\n'.join(code_lines)
                    
                    # Create paragraph for code block
                    p = doc.add_paragraph()
                    p_format = p.paragraph_format
                    p_format.left_indent = Inches(0.25)
                    p_format.right_indent = Inches(0.25)
                    p_format.space_before = Pt(6)
                    p_format.space_after = Pt(6)
                    
                    # Add border/box
                    add_border_to_paragraph(p)
                    
                    # Add code text
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Add spacing after code block
                    p_after = doc.add_paragraph()
                    p_after.paragraph_format.space_after = Pt(6)
                
                code_lines = []
                continue
        
        if in_code_block:
            # Collect code lines
            code_lines.append(stripped)
            continue
        
        # Empty line
        if not stripped:
            if i > 0 and lines[i-1].strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
            continue
        
        # Headers
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            p_format = p.paragraph_format
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Arial'
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            p_format = p.paragraph_format
            p_format.space_before = Pt(10)
            p_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Arial'
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            p_format = p.paragraph_format
            p_format.space_before = Pt(8)
            p_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
        elif stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:].strip(), level=4)
            p_format = p.paragraph_format
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = 'Arial'
        
        # Horizontal rule
        elif stripped.startswith('---'):
            p = doc.add_paragraph('_' * 60)
            p_format = p.paragraph_format
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(12)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(200, 200, 200)
        
        # Bullet points
        elif stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.25)
            p_format.space_after = Pt(3)
            add_formatted_run(p, text)
        
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_after = Pt(6)
            p_format.first_line_indent = Inches(0) if stripped.startswith(('A.', 'B.', 'C.', 'D.', 'E.')) else Inches(0)
            
            # Special formatting for multiple choice options
            if re.match(r'^[A-E]\.\s', stripped):
                p_format.left_indent = Inches(0.3)
                p_format.first_line_indent = Inches(-0.3)
            
            add_formatted_run(p, stripped)
    
    # Save document
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Successfully converted {md_path.name} to {docx_path.name}")
    return docx_path


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='Convert markdown to .docx with enhanced formatting')
    parser.add_argument('input', nargs='?', help='Input markdown file')
    parser.add_argument('-o', '--output', help='Output .docx file')
    parser.add_argument('--batch', action='store_true', help='Convert all .md files in question_sets folder')
    
    args = parser.parse_args()
    
    if args.batch:
        # Convert all markdown files in question_sets folder
        script_dir = Path(__file__).parent
        midterm_dir = script_dir.parent
        question_sets_dir = midterm_dir / 'question_sets'
        docx_dir = midterm_dir / 'docx'
        
        if not question_sets_dir.exists():
            print(f"Error: {question_sets_dir} does not exist")
            sys.exit(1)
        
        md_files = list(question_sets_dir.glob('*.md'))
        
        if not md_files:
            print(f"No .md files found in {question_sets_dir}")
            sys.exit(1)
        
        print(f"Found {len(md_files)} markdown file(s) to convert...")
        
        for md_file in md_files:
            output_file = docx_dir / md_file.with_suffix('.docx').name
            try:
                convert_markdown_simple(md_file, output_file)
            except Exception as e:
                print(f"Error converting {md_file.name}: {e}")
                import traceback
                traceback.print_exc()
        
        print(f"\nAll conversions complete! Files saved to: {docx_dir}")
    
    elif args.input:
        input_file = Path(args.input)
        output_file = Path(args.output) if args.output else None
        convert_markdown_simple(input_file, output_file)
    else:
        # Default: convert all files in question_sets
        script_dir = Path(__file__).parent
        midterm_dir = script_dir.parent
        question_sets_dir = midterm_dir / 'question_sets'
        docx_dir = midterm_dir / 'docx'
        
        md_files = list(question_sets_dir.glob('*.md'))
        
        if md_files:
            print(f"Converting {len(md_files)} file(s)...")
            for md_file in md_files:
                output_file = docx_dir / md_file.with_suffix('.docx').name
                convert_markdown_simple(md_file, output_file)
            print(f"\nAll conversions complete! Files saved to: {docx_dir}")
        else:
            print("No markdown files found. Usage:")
            print("  python markdown_to_docx.py <input.md>")
            print("  python markdown_to_docx.py --batch")
