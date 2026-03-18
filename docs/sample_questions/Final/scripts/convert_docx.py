"""
Script to convert .docx file to readable markdown format
"""
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("python-docx library is required. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document


def docx_to_markdown(docx_path, output_path=None):
    """
    Convert a .docx file to markdown format

    Args:
        docx_path: Path to the .docx file
        output_path: Path for output file (optional, defaults to same name with .md extension)
    """
    docx_path = Path(docx_path)

    if not docx_path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    if output_path is None:
        output_path = docx_path.with_suffix('.md')
    else:
        output_path = Path(output_path)

    doc = Document(docx_path)
    markdown_lines = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.replace('Heading ', '')
                try:
                    level_num = int(level)
                    markdown_lines.append('#' * level_num + ' ' + text)
                except ValueError:
                    markdown_lines.append('## ' + text)
            else:
                is_bold = any(run.bold for run in paragraph.runs)
                if is_bold:
                    markdown_lines.append('**' + text + '**')
                else:
                    markdown_lines.append(text)
            markdown_lines.append('')

    for table in doc.tables:
        markdown_lines.append('')
        for i, row in enumerate(table.rows):
            row_text = '| ' + ' | '.join([cell.text.strip() for cell in row.cells]) + ' |'
            markdown_lines.append(row_text)
            if i == 0:
                markdown_lines.append('| ' + ' | '.join(['---'] * len(row.cells)) + ' |')
        markdown_lines.append('')

    output_path.write_text('\n'.join(markdown_lines), encoding='utf-8')
    print(f"Successfully converted {docx_path.name} to {output_path.name}")
    return output_path


if __name__ == "__main__":
    script_dir = Path(__file__).parent
    final_dir = script_dir.parent
    original_docx = final_dir / "Orginal" / "ABT182_sample_questions_for_final.docx"
    output_md = final_dir / "original" / "ABT182_sample_questions_for_final.md"

    if len(sys.argv) > 1:
        original_docx = Path(sys.argv[1])
        output_md = None

    output_file = docx_to_markdown(original_docx, output_md)
    print(f"\nOutput saved to: {output_file}")
